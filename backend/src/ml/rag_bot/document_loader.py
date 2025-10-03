"""
Document Loader - загрузка и чтение различных типов документов
"""
from typing import List, Dict, Any, Optional
from pathlib import Path
import mimetypes
from dataclasses import dataclass

@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]
    source: str

class DocumentLoader:
    def __init__(self):
        print("DocumentLoader initialized")
    
    def load_file(self, file_path: str) -> Document:
        """
        Загрузить один файл
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Document объект с содержимым и метаданными
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
            
        # Определяем тип файла
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = path.suffix.lower()
        
        print(f"Загружаем файл: {path.name} ({file_extension})")
        
        # Выбираем метод загрузки в зависимости от типа файла
        if file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css']:
            content = self._load_text_file(path)
        elif file_extension == '.pdf':
            content = self._load_pdf_file(path)
        elif file_extension in ['.docx', '.doc']:
            content = self._load_docx_file(path)
        else:
            # Пробуем как текстовый файл
            try:
                content = self._load_text_file(path)
            except UnicodeDecodeError:
                raise ValueError(f"Неподдерживаемый тип файла: {file_extension}")
        
        # Создаем метаданные
        metadata = {
            "filename": path.name,
            "file_path": str(path.absolute()),
            "file_size": path.stat().st_size,
            "file_extension": file_extension,
            "mime_type": mime_type,
            "content_length": len(content)
        }
        
        return Document(
            content=content,
            metadata=metadata,
            source=str(path.absolute())
        )
    
    def load_directory(self, directory_path: str, extensions: Optional[List[str]] = None) -> List[Document]:
        """
        Загрузить все файлы из директории
        
        Args:
            directory_path: Путь к директории
            extensions: Список расширений для фильтрации (например, ['.txt', '.pdf'])
            
        Returns:
            Список Document объектов
        """
        path = Path(directory_path)

        if not path.exists() or not path.is_dir():
            raise ValueError(f"Директория не найдена: {directory_path}")
        
        if extensions is None:
            extensions = ['.txt', '.md', '.pdf', '.docx', '.doc', '.py', '.js', '.html']
        
        documents = []
        for file_path in path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    document = self.load_file(str(file_path))
                    documents.append(document)
                except Exception as e:
                    print(f"Ошибка при загрузке файла {file_path}: {e}")
                    continue

        print(f"Загружено {len(documents)} документов из {directory_path}")
        return documents
    
    def _load_text_file(self, path: Path) -> str:
        encodings = ['utf-8', 'cp1251', 'latin1']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError(f"Не удалось декодировать файл {path} с кодировками {encodings}")
        
    def _load_pdf_file(self, path: Path) -> str:
        """Загрузить PDF файл"""
        try:
            import PyPDF2
            
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                
                return content.strip()
                
        except ImportError:
            raise ImportError("Для работы с PDF установите: pip install PyPDF2")
        except Exception as e:
            raise ValueError(f"Ошибка чтения PDF: {e}")
    
    def _load_docx_file(self, path: Path) -> str:
        """Загрузить DOCX файл"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            return content.strip()
            
        except ImportError:
            raise ImportError("Для работы с DOCX установите: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Ошибка чтения DOCX: {e}")
    def load_url(self, url: str) -> Document:
        """
        Загрузить содержимое веб-страницы
        
        Args:
            url: URL веб-страницы
            
        Returns:
            Document объект с содержимым страницы
        """
        try:
            import requests
            from bs4 import BeautifulSoup
            
            response = requests.get(url)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            content = soup.get_text()
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            content = soup.get_text()

            lines = (line.strip() for line in content.splitlines())
            content = '\n'.join(line for line in lines if line)
             
            metadata = {
                "url": url,
                "title": soup.title.string if soup.title else "Без заголовка",
                "content_length": len(content),
                "status_code": response.status_code,
            }

            return Document(
                content=content,
                metadata=metadata,
                source=url
            )
        except ImportError:
            raise ImportError("Для работы с URL установите: pip install requests beautifulsoup4")
        except Exception as e:
            raise ValueError(f"Ошибка загрузки URL: {e}")